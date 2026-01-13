#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
import re

# 打开文档
doc_path = "/Users/kexiaobin/Downloads/数字工厂V2.0优化版.docx"
doc = Document(doc_path)

print("=== 分析文档内容（前100个段落）===\n")

# 显示前100个段落的内容
for i, paragraph in enumerate(doc.paragraphs[:100]):
    text = paragraph.text.strip()

    if text:
        # 显示以数字、括号或特殊符号开头的内容
        first_chars = text[:10] if len(text) >= 10 else text

        # 检查是否包含编号模式
        has_numbering = (
            re.match(r'^\d+\.', text) or  # 1. 2. 3.
            re.match(r'^\d+\.\d+\.', text) or  # 1.1 1.2
            re.match(r'^\d+\.\d+\.\d+', text) or  # 1.1.1
            re.match(r'^[（(]\d+[)）]', text) or  # (1)
            re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', text)  # 圆圈数字
        )

        if has_numbering or i < 30:  # 显示前30个，或包含编号的
            style = paragraph.style.name if paragraph.style else 'Normal'
            print(f"{i+1:3d}. [{style:20s}] {text[:100]}")
            print()

    if i >= 99:
        break

# 统计编号类型
print("\n=== 统计编号类型 ===\n")

pattern_1_1 = 0
pattern_1_1_1 = 0
pattern_parenthesis = 0
pattern_circle = 0

for paragraph in doc.paragraphs:
    text = paragraph.text.strip()

    if re.match(r'^\d+\.\d+\s', text):
        pattern_1_1 += 1
    elif re.match(r'^\d+\.\d+\.\d+\s', text):
        pattern_1_1_1 += 1
    elif re.match(r'^[（(]\d+[)）]\s', text):
        pattern_parenthesis += 1
    elif re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]\s', text):
        pattern_circle += 1

print(f"'1.1' 格式: {pattern_1_1} 个")
print(f"'1.1.1' 格式: {pattern_1_1_1} 个")
print(f"'(1)' 格式: {pattern_parenthesis} 个")
print(f"'①' 格式: {pattern_circle} 个")
