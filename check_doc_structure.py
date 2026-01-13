#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.oxml.ns import qn

# 打开文档
doc_path = "/Users/kexiaobin/Downloads/数字工厂V2.0优化版.docx"
doc = Document(doc_path)

print("=== 文档结构分析 ===\n")

print(f"总段落数: {len(doc.paragraphs)}")
print(f"总表格数: {len(doc.tables)}")
print(f"总节数: {len(doc.sections)}")

# 检查body元素
body = doc._element.body
print(f"\nBody直接子元素数: {len(body)}")

# 遍历body的所有子元素
for i, elem in enumerate(body[:20]):
    tag_name = elem.tag
    namespace = elem.nsmap.get(elem.prefix, '')
    print(f"{i+1}. 标签: {tag_name}")

    # 尝试获取文本内容
    if hasattr(elem, 'text'):
        if elem.text and elem.text.strip():
            print(f"   文本: {elem.text[:100]}...")
