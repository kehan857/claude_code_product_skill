#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
import re

def set_cell_font(cell, font_name='宋体', font_size=9):
    """设置表格单元格内所有文本的字体为宋体小5号（9磅）"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)

def remove_numbering(paragraph):
    """移除段落的编号"""
    pPr = paragraph._element.pPr
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)

def add_custom_bullet_with_text(paragraph, prefix, font_name='宋体', font_size=10.5):
    """移除编号并添加自定义文本前缀"""
    # 移除现有编号
    remove_numbering(paragraph)

    # 保存原文本
    original_text = paragraph.text

    # 清除段落内容
    for run in paragraph.runs:
        run.text = ''

    # 添加新的run
    run = paragraph.add_run(f'{prefix} {original_text}')
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)

# 打开文档
doc_path = "/Users/kexiaobin/Downloads/数字工厂概要设计V2.0版.docx"
doc = Document(doc_path)

print("开始处理文档...\n")

# 1. 处理所有表格：设置字体为宋体小5号（9磅）
print("1. 处理表格字体...")
table_count = 0
for table in doc.tables:
    table_count += 1
    for row in table.rows:
        for cell in row.cells:
            set_cell_font(cell, font_name='宋体', font_size=9)

print(f"✓ 已处理 {table_count} 个表格的字体（宋体 小5号）\n")

# 2. 处理项目符号和编号
print("2. 处理项目符号和编号...")

# 圆圈数字符号
circle_nums = ['①', '②', '③', '④', '⑤', '⑥', '⑦', '⑧', '⑨', '⑩',
               '⑪', '⑫', '⑬', '⑭', '⑮', '⑯', '⑰', '⑱', '⑲', '⑳']

parenthesis_count = 0
circle_count = 0
modified_count = 0

# 收集有编号的段落信息
numbered_paragraphs = []
for i, paragraph in enumerate(doc.paragraphs):
    pPr = paragraph._element.pPr
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            ilvl_elem = numPr.find(qn('w:ilvl'))
            numId_elem = numPr.find(qn('w:numId'))

            level = 0
            num_id = 1

            if ilvl_elem is not None:
                level = int(ilvl_elem.get(qn('w:val')))
            if numId_elem is not None:
                num_id = int(numId_elem.get(qn('w:val')))

            text = paragraph.text.strip()
            if text:
                numbered_paragraphs.append({
                    'index': i,
                    'paragraph': paragraph,
                    'level': level,
                    'num_id': num_id,
                    'text': text
                })

print(f"  找到 {len(numbered_paragraphs)} 个带编号的段落")

# 处理编号段落（需要反向处理，避免索引变化）
for item in reversed(numbered_paragraphs):
    paragraph = item['paragraph']
    level = item['level']
    text = item['text']

    # 根据层级和编号ID决定处理方式
    # level 0 且 numId 1-5 通常是主要列表（黑点）
    # level 1 通常是次要列表（空心点）

    if level == 0:
        # 第一层：改为括号 (1)(2)(3)...
        parenthesis_count += 1
        add_custom_bullet_with_text(paragraph, f'({parenthesis_count})')
        modified_count += 1

    elif level == 1:
        # 第二层：改为圆圈数字 ①②③...
        circle_index = circle_count % len(circle_nums)
        circle_symbol = circle_nums[circle_index]
        circle_count += 1
        add_custom_bullet_with_text(paragraph, circle_symbol)
        modified_count += 1

print(f"✓ 已修改 {modified_count} 个项目符号")
print(f"  - 黑点 -> 括号：{parenthesis_count} 处")
print(f"  - 空心点 -> 圆圈：{circle_count} 处")

# 保存文档
output_path = "/Users/kexiaobin/Desktop/其他/claude code/数字工厂概要设计V2.0版-已修改.docx"
doc.save(output_path)

print(f"\n✅ 文档处理完成！")
print(f"📄 保存位置：{output_path}")
print("\n修改内容：")
print("  ✓ 所有表格内字体调整为：宋体 小5号（9磅）")
print("  ✓ 第一层编号（黑点）改为：括号符号 (1)(2)(3)...")
print("  ✓ 第二层编号（空心点）改为：圆圈符号 ①②③...")
