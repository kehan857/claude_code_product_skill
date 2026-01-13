#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 打开文档
doc_path = "/Users/kexiaobin/Downloads/数字工厂设计V2.0.docx"
doc = Document(doc_path)

print("开始修改BOM展开示例...\n")

# 新的BOM展开示例内容
new_bom_example = """产品A（1套）
├─ 物料B（2个）
│  ├─ 物料D（5个）
│  └─ 物料E（3个）
├─ 物料C（1套）
│  ├─ 物料F（10个）
│  └─ 物料G（2个）
└─ 物料H（1个）- 外购物料

如果MPS计划生产100套A，则需要：
• 物料B：200个
• 物料C：100套
• 物料D：1000个（200个B × 5）
• 物料E：600个（200个B × 3）
• 物料F：1000个（100套C × 10）
• 物料G：200个（100套C × 2）
• 物料H：100个"""

modified_count = 0

# 修改段落237（原BOM树结构）
for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()

    # 识别BOM树结构的段落
    if "产品A（1套）" in text and "物料B（2个）" in text and "├─" in text:
        print(f"找到BOM树结构段落 {i+1}")

        # 清除原有内容
        for run in paragraph.runs:
            run.text = ''

        # 添加新内容（使用等宽字体以保持树形结构）
        run = paragraph.add_run(new_bom_example)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')

        modified_count += 1
        print(f"  ✓ 已修改BOM树结构")

    # 识别MRP计算段落
    elif "如果MPS计划生产100套A" in text and "物料B：200个" in text:
        print(f"找到MRP计算段落 {i+1}")

        # 这个段落已经在新的BOM示例中了，清除它
        for run in paragraph.runs:
            run.text = ''

        modified_count += 1
        print(f"  ✓ 已清除旧的MRP计算段落（已整合到新的示例中）")

# 保存文档
output_path = "/Users/kexiaobin/Desktop/其他/claude code/数字工厂设计V2.0-已修改.docx"
doc.save(output_path)

print(f"\n✅ 文档修改完成！")
print(f"📄 保存位置：{output_path}")
print(f"\n修改了 {modified_count} 处内容")
print("\n新的BOM展开示例：")
print(new_bom_example)
