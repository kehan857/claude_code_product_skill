#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.oxml.ns import qn
import re

# 打开文档
doc_path = "/Users/kexiaobin/Downloads/数字工厂V2.0优化版.docx"
doc = Document(doc_path)

print("=== 分析文档编号结构 ===\n")

# 收集不同层级的编号样式
numbering_samples = {
    'level_0': [],  # 1.1, 1.2 等
    'level_1': [],  # 1.1.1, 1.1.2 等
    'level_2': [],  # 项目符号
    'headings': []
}

for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()

    if not text or len(text) < 3:
        continue

    # 检测不同类型的编号
    # 1.1 格式
    if re.match(r'^\d+\.\d+\s', text):
        numbering_samples['level_0'].append({
            'index': i,
            'text': text[:80],
            'style': paragraph.style.name if paragraph.style else 'Normal'
        })

    # 1.1.1 格式
    elif re.match(r'^\d+\.\d+\.\d+\s', text):
        numbering_samples['level_1'].append({
            'index': i,
            'text': text[:80],
            'style': paragraph.style.name if paragraph.style else 'Normal'
        })

    # 标题
    elif text.startswith('第') and '章' in text:
        numbering_samples['headings'].append({
            'index': i,
            'text': text[:80]
        })

# 显示分析结果
print(f"找到 {len(numbering_samples['headings'])} 个章节标题")
for item in numbering_samples['headings'][:5]:
    print(f"  段落{item['index']+1}: {item['text']}...")

print(f"\n找到 {len(numbering_samples['level_0'])} 个 '1.1' 格式的编号")
for item in numbering_samples['level_0'][:10]:
    print(f"  段落{item['index']+1} ({item['style']}): {item['text']}...")

print(f"\n找到 {len(numbering_samples['level_1'])} 个 '1.1.1' 格式的编号")
for item in numbering_samples['level_1'][:10]:
    print(f"  段落{item['index']+1} ({item['style']}): {item['text']}...")

# 检查是否有项目符号
print("\n=== 检查项目符号 ===")
bullet_count = 0
for paragraph in doc.paragraphs:
    pPr = paragraph._element.pPr
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            bullet_count += 1
            if bullet_count <= 5:
                print(f"  段落项目符号: {paragraph.text[:50]}...")

print(f"\n总共找到 {bullet_count} 个使用项目符号的段落")
