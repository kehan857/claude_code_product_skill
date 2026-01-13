#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
import re

# 打开文档
doc_path = "/Users/kexiaobin/Downloads/数字工厂设计V2.0.docx"
doc = Document(doc_path)

print("=== 搜索文档中的BOM、MRP相关内容 ===\n")

keywords = ['BOM', 'MRP', '物料需求', '展开', '示例', '流程']
found_paragraphs = []

for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    if text and any(keyword in text for keyword in keywords):
        found_paragraphs.append({
            'index': i,
            'text': text,
            'style': paragraph.style.name if paragraph.style else 'Normal'
        })

print(f"找到 {len(found_paragraphs)} 个相关段落\n")

# 显示前20个相关段落
for item in found_paragraphs[:20]:
    print(f"段落 {item['index']+1} ({item['style']}):")
    print(f"  {item['text'][:100]}...")
    print()

# 搜索表格中的相关内容
print("\n=== 搜索表格中的BOM、MRP相关内容 ===\n")

for table_idx, table in enumerate(doc.tables[:10]):
    table_text = ""
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                table_text += cell.text.strip() + " | "

    if any(keyword in table_text for keyword in keywords):
        print(f"表格 {table_idx + 1}:")
        print(f"  {table_text[:200]}...")
        print()
