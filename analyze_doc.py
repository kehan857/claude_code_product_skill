#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.oxml.ns import qn
import re

# 打开文档
doc_path = "/Users/kexiaobin/Downloads/数字工厂概要设计V2.0版.docx"
doc = Document(doc_path)

print("=== 分析文档中的项目符号格式 ===\n")

# 检查前50个段落
sample_count = 0
for i, paragraph in enumerate(doc.paragraphs[:50]):
    text = paragraph.text.strip()
    if text:
        sample_count += 1
        print(f"{sample_count}. 段落 {i+1}:")
        print(f"   文本: {text[:50]}...")

        # 检查段落样式
        if paragraph.style:
            print(f"   样式: {paragraph.style.name}")

        # 检查是否有项目符号或编号
        pPr = paragraph._element.pPr
        if pPr is not None:
            # 检查项目符号
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                print(f"   ✓ 包含编号/项目符号")

                # 获取编号信息
                ilvl = numPr.find(qn('w:ilvl'))
                numId = numPr.find(qn('w:numId'))

                if ilvl is not None:
                    print(f"   - 层级: {ilvl.get(qn('w:val'))}")
                if numId is not None:
                    print(f"   - 编号ID: {numId.get(qn('w:val'))}")

        # 检查段落开头的字符
        if text:
            first_char = text[0]
            print(f"   首字符: '{first_char}' (Unicode: U+{ord(first_char):04X})")

        print()

    if sample_count >= 20:
        break

print("\n=== 检查是否有特殊符号 ===")

# 搜索包含常见项目符号的段落
bullet_chars = {
    '\u2022': '实心点 •',
    '\u25E6': '空心点 ◦',
    '\u25CB': '圆圈 ○',
    '\u25CF': '黑色圆圈 ●',
    '\u25A0': '方块 ■',
    '\u25AA': '小方块 ▪',
}

found_bullets = {}
for paragraph in doc.paragraphs:
    text = paragraph.text
    for char, name in bullet_chars.items():
        if char in text:
            if name not in found_bullets:
                found_bullets[name] = []
            found_bullets[name].append(text[:60])

if found_bullets:
    print("\n找到以下项目符号：")
    for name, examples in found_bullets.items():
        print(f"\n{name}:")
        for ex in examples[:3]:
            print(f"  - {ex}...")
else:
    print("\n未在文本中找到常见的项目符号字符")
    print("项目符号可能是通过 Word 的编号功能实现的")
