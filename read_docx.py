#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document

# 读取文档
doc_path = "/Users/kexiaobin/Downloads/tencent_authorization.docx"
doc = Document(doc_path)

# 打印所有段落内容
print("=== 文档内容 ===")
for i, paragraph in enumerate(doc.paragraphs, 1):
    if paragraph.text.strip():
        print(f"{i}. {paragraph.text}")

# 打印所有表格
print("\n=== 表格内容 ===")
for table_idx, table in enumerate(doc.tables, 1):
    print(f"\n表格 {table_idx}:")
    for row_idx, row in enumerate(table.rows):
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip())
        print(f"  行{row_idx+1}: {' | '.join(row_data)}")
