#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document

# 打开文档
doc_path = "/Users/kexiaobin/Downloads/数字工厂V2.0优化版.docx"
doc = Document(doc_path)

print("=== 文档内容（前50段）===\n")

count = 0
for i, paragraph in enumerate(doc.paragraphs):
    if paragraph.text.strip():
        count += 1
        print(f"{count}. {paragraph.text}")
        print()

    if count >= 50:
        break

print(f"\n文档总段落数: {len(doc.paragraphs)}")
