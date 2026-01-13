#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 打开原文档
doc_path = "/Users/kexiaobin/Downloads/tencent_authorization.docx"
doc = Document(doc_path)

# 提供的信息
company_info = {
    "company_full_name": "熵变智元(北京)科技有限公司",
    "credit_code": "91110117MAC9D5925A",
    "penguin_account_name": "熵变AI实验室",
    "account_intro": "想在AI这个高度智能的新维度纪元，帮你在混乱、无序的初期建立有序、减少熵增。",
    "address": "北京市平谷区夏各庄镇马各庄南街83号23068",
    "department": "产品营销部",
    "operator_name": "可晓斌",
    "id_card": "410402199007155833"
}

# 修改段落内容
for paragraph in doc.paragraphs:
    text = paragraph.text
    if "企鹅号名称为：XXX" in text:
        # 替换企鹅号名称和运营者姓名
        new_text = text.replace("XXX", company_info["penguin_account_name"])
        new_text = new_text.replace("授权XXX", f"授权{company_info['operator_name']}")
        paragraph.clear()
        paragraph.add_run(new_text)
        print(f"✓ 已更新授权声明")

# 修改表格内容
if doc.tables:
    table = doc.tables[0]

    # 行1: 企业（单位）全称
    table.rows[0].cells[1].text = company_info["company_full_name"]

    # 行2: 单位所在地
    table.rows[1].cells[1].text = company_info["address"]

    # 行3: 企鹅号名称和账号介绍
    table.rows[2].cells[1].text = company_info["penguin_account_name"]
    table.rows[2].cells[3].text = company_info["account_intro"]

    # 行5: 运营者姓名和身份证号
    table.rows[4].cells[1].text = company_info["operator_name"]
    table.rows[4].cells[3].text = company_info["id_card"]

    # 行6: 所在部门
    table.rows[5].cells[1].text = company_info["department"]

    print("✓ 已更新表格信息")

# 保存新文档
output_path = "/Users/kexiaobin/Desktop/其他/claude code/腾讯授权书-熵变智元.docx"
doc.save(output_path)

print(f"\n✅ 文档已完成并保存到：{output_path}")
print("\n已填写的信息：")
print(f"  • 企业全称：{company_info['company_full_name']}")
print(f"  • 统一社会信用代码：{company_info['credit_code']}")
print(f"  • 企鹅号名称：{company_info['penguin_account_name']}")
print(f"  • 账号介绍：{company_info['account_intro']}")
print(f"  • 注册地址：{company_info['address']}")
print(f"  • 部门：{company_info['department']}")
print(f"  • 运营者：{company_info['operator_name']}")
print(f"  • 身份证号：{company_info['id_card']}")
