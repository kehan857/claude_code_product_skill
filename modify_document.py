#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import copy

def set_cell_font(cell, font_name='宋体', font_size=9):
    """设置表格单元格内所有文本的字体为宋体小5号（9磅）"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)

def set_paragraph_font(paragraph, font_name='宋体', font_size=10.5):
    """设置段落字体"""
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)

def detect_bullet_type(text):
    """检测文本开头的项目符号类型"""
    if not text.strip():
        return None, 0

    patterns = [
        (r'^[\u2022]\s*', 'bullet', 1),           # 实心点 •
        (r'^[\u25CF]\s*', 'black_circle', 1),     # 黑色圆圈 ●
        (r'^[\u25E6]\s*', 'hollow', 2),           # 空心点 ◦
        (r'^[\u25CB]\s*', 'circle', 2),           # 圆圈 ○
        (r'^\(\d+\)\s*', 'parenthesis', 1),       # 已经是括号 (1)
        (r'^[①②③④⑤⑥⑦⑧⑨⑩]\s*', 'circle_num', 2), # 已经是圆圈数字
    ]

    for pattern, bullet_type, level in patterns:
        if re.match(pattern, text):
            return bullet_type, level

    return None, 0

def get_next_parenthesis_num(current_num):
    """获取下一个括号数字"""
    return current_num + 1

def get_next_circle_num(current_num):
    """获取下一个圆圈数字"""
    circle_nums = ['①', '②', '③', '④', '⑤', '⑥', '⑦', '⑧', '⑨', '⑩',
                   '⑪', '⑫', '⑬', '⑭', '⑮', '⑯', '⑰', '⑱', '⑲', '⑳']
    if current_num < len(circle_nums):
        return circle_nums[current_num]
    return circle_nums[-1]

def remove_bullet_prefix(text):
    """移除项目符号前缀"""
    # 移除所有已知的项目符号
    text = re.sub(r'^[\u2022\u25E6\u25CB\u25CF\u25A0\u25AA\-\*]\s*', '', text)
    return text.strip()

# 打开文档
doc_path = "/Users/kexiaobin/Downloads/数字工厂概要设计V2.0版.docx"
doc = Document(doc_path)

print("开始处理文档...")

# 1. 处理所有表格：设置字体为宋体小5号（9磅）
table_count = 0
for table in doc.tables:
    table_count += 1
    for row in table.rows:
        for cell in row.cells:
            set_cell_font(cell, font_name='宋体', font_size=9)

print(f"✓ 已处理 {table_count} 个表格的字体（宋体 小5号）")

# 2. 处理项目符号
paragraph_count = 0
modified_count = 0
parenthesis_counter = 0
circle_counter = 0

for paragraph in doc.paragraphs:
    if paragraph.text.strip():
        paragraph_count += 1
        text = paragraph.text
        bullet_type, level = detect_bullet_type(text)

        if bullet_type in ['bullet', 'black_circle']:
            # 黑点 -> 改为括号 (1)(2)(3)...
            modified_count += 1
            parenthesis_counter = get_next_parenthesis_num(parenthesis_counter)

            # 清除原有文本
            for run in paragraph.runs:
                run.text = ''

            # 添加新文本
            new_text = remove_bullet_prefix(text)
            run = paragraph.add_run(f'({parenthesis_counter}) {new_text}')
            set_paragraph_font(run, font_name='宋体', font_size=10.5)

        elif bullet_type in ['hollow', 'circle']:
            # 空心点 -> 改为圆圈符号 ①②③...
            modified_count += 1
            circle_index = circle_counter % 20  # 支持20个圆圈数字
            circle_symbol = get_next_circle_num(circle_index)
            circle_counter += 1

            # 清除原有文本
            for run in paragraph.runs:
                run.text = ''

            # 添加新文本
            new_text = remove_bullet_prefix(text)
            run = paragraph.add_run(f'{circle_symbol} {new_text}')
            set_paragraph_font(run, font_name='宋体', font_size=10.5)

print(f"✓ 已扫描 {paragraph_count} 个段落")
print(f"✓ 已修改 {modified_count} 个项目符号")
print(f"  - 黑点 -> 括号：{parenthesis_counter} 处")
print(f"  - 空心点 -> 圆圈：{circle_counter} 处")

# 保存文档
output_path = "/Users/kexiaobin/Desktop/其他/claude code/数字工厂概要设计V2.0版-已修改.docx"
doc.save(output_path)

print(f"\n✅ 文档处理完成！")
print(f"📄 保存位置：{output_path}")
print("\n修改内容：")
print("  ✓ 所有表格内字体调整为：宋体 小5号（9磅）")
print("  ✓ 黑点（•、●）改为：括号符号 (1)(2)(3)...")
print("  ✓ 空心点（◦、○）改为：圆圈符号 ①②③...")
