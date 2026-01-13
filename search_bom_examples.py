#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
import re

# 打开文档
doc_path = "/Users/kexiaobin/Downloads/数字工厂设计V2.0.docx"
doc = Document(doc_path)

print("=== 搜索MRP、BOM展开相关内容 ===\n")

# 搜索包含更具体关键词的段落
specific_keywords = ['MPS', 'MRP', '需求计算', '展开', '套', '个', '物料']
found_paragraphs = []

for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    if text and len(text) > 20:  # 只关注较长的段落
        # 检查是否包含多个相关关键词
        keyword_count = sum(1 for kw in specific_keywords if kw in text)
        if keyword_count >= 2:
            found_paragraphs.append({
                'index': i,
                'text': text,
                'keyword_count': keyword_count
            })

# 按关键词数量排序
found_paragraphs.sort(key=lambda x: x['keyword_count'], reverse=True)

print(f"找到 {len(found_paragraphs)} 个可能包含BOM展开示例的段落\n")

# 显示最相关的段落
for item in found_paragraphs[:15]:
    print(f"段落 {item['index']+1} (关键词数: {item['keyword_count']}):")
    print(f"  {item['text'][:150]}...")
    print()

# 搜索表格中的具体示例
print("\n=== 搜索表格中的详细示例 ===\n")

for table_idx, table in enumerate(doc.tables):
    has_example = False
    table_content = []

    for row_idx, row in enumerate(table.rows):
        row_text = []
        for cell in row.cells:
            row_text.append(cell.text.strip())

        row_str = " | ".join(row_text)
        if any(kw in row_str for kw in ['示例', '产品', '物料', '套', '个']):
            has_example = True
            table_content.append(f"  行{row_idx+1}: {row_str[:120]}")

    if has_example:
        print(f"表格 {table_idx + 1} 可能包含示例:")
        for line in table_content[:5]:
            print(line)
        print()
